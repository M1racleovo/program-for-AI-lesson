"""生成进度报告 Word 文档"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def make_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ========== 封面标题 ==========
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《人工智能导论》课程项目\n进 展 报 告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()
    doc.add_paragraph()

    # 项目信息表
    info_items = [
        ("项  目  题  目：", "FaceFair：面向多人种与多属性的人脸识别公平性评估与偏见分析系统"),
        ("所  在  学  院：", "计算机学院"),
        ("项  目  组  长：", "余海阳"),
        ("小  组  成  员：", "余海阳、马钰轩、刘星悦"),
        ("开  题  时  间：", "2026年 4月 12日"),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(label)
        run_l.font.size = Pt(14)
        run_l.font.name = "宋体"
        run_l.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run_v = p.add_run(value)
        run_v.font.size = Pt(14)
        run_v.font.name = "宋体"
        run_v.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if label == "项  目  题  目：":
            run_v.bold = True

    doc.add_page_break()

    # ========== 正文 ==========
    section_title_font_size = Pt(16)
    body_font_size = Pt(12)

    def add_section(title_text, content_text):
        # 标题
        h = doc.add_heading(title_text, level=1)
        for run in h.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = section_title_font_size
            run.font.color.rgb = RGBColor(0, 0, 0)

        # 内容
        for para in content_text.strip().split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            if para.startswith("  - ") or para.startswith("  • "):
                p.paragraph_format.left_indent = Cm(1.5)
                text = para.lstrip("  -• ")
            else:
                text = para.strip()
            run = p.add_run(text)
            run.font.size = body_font_size
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---- 一、项目任务 ----
    add_section("一、项目任务", """
本项目旨在构建一个名为 FaceFair 的人脸识别公平性评估与偏见分析系统。
随着深度学习技术在人脸识别领域的广泛应用，算法在不同人种、性别和年龄群体间的公平性问题日益凸显。
本项目基于 FairFace 数据集和 DeepFace 深度学习框架，对主流预训练模型（如 VGG-Face、Facenet 等）在不同肤色人群上的识别准确率进行量化分析，
评估性别和年龄因素如何与种族因素共同影响识别结果，并通过可视化技术直观展示识别偏差，探讨偏见产生的潜在原因。
研究内容涵盖自动化评估管线的搭建、多维度属性标注对比，以及公平性报告的生成。
""")

    # ---- 二、技术方案 ----
    add_section("二、技术方案", """
本项目采用 Python 作为核心开发语言，结合以下技术栈构建自动化分析管线：

1. 人脸分析框架：使用 DeepFace 库作为核心分析引擎。DeepFace 是一个轻量级的人脸识别与分析框架，
   内部集成了 VGG-Face、Facenet、ArcFace 等多种工业级预训练模型，支持对人脸图像进行年龄、性别、种族和表情等属性的自动预测。
   通过封装统一的 API 接口，可实现对成规模数据集的快速批量分析。

2. 深度学习后端：使用 TensorFlow/Keras 作为底层深度学习推理引擎，通过 tf-keras 兼容层确保 DeepFace 与最新 TensorFlow 版本的兼容性。

3. 人脸检测：使用 OpenCV 作为默认的人脸检测后端，支持 retinaface、mtcnn、dlib、ssd 等多种检测器的灵活切换。

4. 数据处理：使用 Pandas 进行分析结果的整理与统计分析，NumPy 提供高效的数值计算支持，OpenCV 和 Pillow 用于图像的预处理与加载。

5. 可视化：使用 Matplotlib 和 Seaborn 生成各种族准确率柱状图、种族×性别热力图、混淆矩阵、置信度箱线图和公平性雷达图等可视化图表。

6. AI 辅助：通过 Ollama 平台运行 Gemma 2b 本地大语言模型，辅助理解报错信息、优化 Python 代码逻辑，以及润色论文中的社会伦理维度分析。

技术路线：FairFace 数据集 → 人脸检测与对齐 → DeepFace 多属性预测 → 结果收集与匹对 → CSV 数据输出 → 统计分析 → 可视化报告。
""")

    # ---- 三、实验方案 ----
    add_section("三、实验方案", """
第3-4周的实验方案聚焦于小规模端到端验证，确保整个分析管线能够正常运转。具体步骤如下：

1. 环境搭建（第3周）：
   • 安装 Python 3.8+ 并配置虚拟环境
   • 通过 pip 安装 deepface、tensorflow、tf-keras、pandas、numpy、opencv-python、matplotlib、seaborn、tqdm 等依赖库
   • 验证 DeepFace 能够正确加载预训练模型（Facenet），并完成基本的属性分析功能
   • 运行 run_test.py 测试脚本，检查所有模块的导入情况

2. 小规模数据测试（第3-4周）：
   • 准备 30 张以内的测试图像，覆盖不同种族和性别类别
   • 执行 face_fair_analyzer.py 小规模模式（--limit 30），对每张图像进行年龄、性别、种族、表情的四维属性分析
   • 验证分析结果能否正确写入 analysis_results.csv 文件，包含图像路径、真实标签、预测结果、匹对状态等字段
   • 生成分析摘要 summary.txt，统计各人群的识别准确率和偏差情况

3. 可视化验证（第4周）：
   • 运行 make_charts.py 脚本，基于小规模分析结果生成各种族准确率柱状图、混淆矩阵、置信度分布图等
   • 验证图表能够正确渲染并保存为 PNG 文件

4. 管线完整性验证：
   • 确保从图像输入到 CSV 输出再到图表生成的全流程能够自动化串联执行
   • 记录人脸检测失败和属性预测异常的情况，为第5-6周的全量分析积累经验
""")

    # ---- 四、目前进展 ----
    add_section("四、目前进展", """
截至第4周末，本组已完成以下工作：

1. Python 开发环境配置完成：
   • 已安装并配置 Python 3.8+ 运行环境
   • 成功安装 deepface、tensorflow（2.15+）、tf-keras、pandas、numpy、opencv-python、matplotlib、seaborn、tqdm 等全部依赖库
   • 通过 run_test.py 脚本验证了所有模块的导入和基础功能

2. 项目代码框架搭建完成（位于 第7组项目/ 目录下）：
   • config.py —— 全局配置文件，定义了路径、模型参数、数据集类别映射、样本量配置等
   • face_fair_analyzer.py —— 核心分析脚本（约250行），实现了数据集加载、DeepFace 批量分析、种族/性别匹对逻辑、CSV 结果输出和摘要生成
   • make_charts.py —— 可视化脚本（约200行），支持生成5种图表：种族准确率柱状图、种族×性别热力图、混淆矩阵、置信度箱线图、公平性雷达图
   • run_test.py —— 环境与功能测试脚本，覆盖 Python 版本检查、依赖导入检查、DeepFace 功能验证、CSV 读写测试
   • requirements.txt —— 依赖清单文件

3. 小规模功能验证完成：
   • DeepFace 人脸分析功能经测试可用，能够正常提取年龄、性别、种族、表情四种属性的预测结果
   • 种族比较逻辑（_compare_race 函数）经过 7 组测试用例验证，能够正确处理不同命名规范下的标签匹配
   • CSV 结果输入输出功能正常
   • 可视化图表生成功能就绪，待积累更多数据后即可生成有意义的分析图表

4. 模块导入与集成测试通过：
   • face_fair_analyzer 模块能够被正确导入和调用
   • 配置模块（config.py）中的所有路径和参数均能正确读取
""")

    # ---- 五、待解决的问题 ----
    add_section("五、待解决的问题", """
以下问题将在第5-8周逐步解决：

1. 全量数据集分析（第5-6周）：
   • 当前仅完成了小规模（30 样本以内）的功能验证，需要将 FairFace 数据集完整解压至 data/fairface/ 目录
   • 执行 python face_fair_analyzer.py --full 命令，进行 1000+ 样本的全量属性分析
   • 完成 analysis_results.csv 的完整数据积累，确保覆盖 7 个种族类别的均衡样本
   • 处理全量分析中可能出现的性能问题（推理速度、内存占用）和批量错误处理

2. 可视化报告生成（第7周）：
   • 基于全量分析数据运行 make_charts.py，生成最终的5张可视化图表
   • 利用 Ollama 平台接入 Gemma 2b 本地大语言模型，辅助撰写"偏见原因分析"章节
   • 整合实验数据与可视化结果，形成完整的分析报告

3. 项目结题（第8周）：
   • 撰写结题论文，总结实验发现和结论
   • 制作项目答辩演示 PPT
   • 准备课程答辩

4. 技术风险与应对：
   • DeepFace 模型下载可能受网络环境影响——建议预先下载模型文件至本地 .deepface/ 目录
   • FairFace 数据集规模较大（约 100k+ 图像）——全量分析前需确认磁盘空间充足
   • 全量推理耗时可能较长——可考虑使用 GPU 加速或分批次执行
""")

    # ---- 六、参考文献 ----
    add_section("六、参考文献", """
[1] Eidinger, E., Enbar, R., & Hassner, T. (2014). Age and gender estimation of unfiltered faces. IEEE Transactions on Information Forensics and Security, 9(12), 2170-2179.

[2] NIST Research Report. (2019). Face Recognition Vendor Test (FRVT) Part 3: Demographic Effects. National Institute of Standards and Technology.

[3] Karkkainen, K., & Joo, J. (2021). FairFace: Face Attribute Dataset for Balanced Race, Gender, and Age for Bias Measurement and Mitigation. In Proceedings of the IEEE/CVF Winter Conference on Applications of Computer Vision (pp. 1548-1558).

[4] Buolamwini, J., & Gebru, T. (2018). Gender Shades: Intersectional Accuracy Disparities in Commercial Gender Classification. In Proceedings of the 1st Conference on Fairness, Accountability and Transparency (pp. 77-91).

[5] Serengil, S. I., & Ozpinar, A. (2020). DeepFace: A Lightweight Deep Face Recognition Library for Python. GitHub Repository.

[6] 《人工智能导论》课程项目指南（2026版）.

[7] Deng, J., Guo, J., Xue, N., & Zafeiriou, S. (2019). ArcFace: Additive Angular Margin Loss for Deep Face Recognition. In Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (pp. 4690-4699).

[8] Schroff, F., Kalenichenko, D., & Philbin, J. (2015). FaceNet: A Unified Embedding for Face Recognition and Clustering. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (pp. 815-823).
""")

    # 保存
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "FaceFair进展报告.docx")
    doc.save(output_path)
    print(f"进度报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    make_report()
